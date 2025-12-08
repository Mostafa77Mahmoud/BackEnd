"""
Document processing service for DOCX manipulation and conversion.
Matches OldStrcturePerfectProject/doc_processing.py exactly.
"""

import os
import uuid
import re
import traceback
import subprocess
import tempfile
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING, WD_BREAK 
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_DIRECTION, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl, CT_TcPr
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
import logging
from flask import current_app

from app.utils.file_helpers import ensure_dir, clean_filename
from app.utils.text_processing import clean_model_response

logger = logging.getLogger(__name__)


def build_structured_text_for_analysis(doc: DocxDocument) -> tuple[str, str]:
    """
    Extracts text from a DOCX document, converting it to a markdown-like format
    that preserves bold, italic, and underline formatting, while also assigning
    unique IDs to paragraphs and table cell content for precise term identification.
    Returns a structured markdown string with IDs and a plain text version.
    """
    structured_markdown = []
    plain_text_parts = []
    para_idx_counter_body = 0
    table_idx_counter_body = 0

    for element in doc.element.body:
        if isinstance(element, CT_P):
            para = Paragraph(element, doc)
            if para.text.strip():
                para_id = f"para_{para_idx_counter_body}"
                
                # Convert paragraph to markdown while preserving formatting
                markdown_line = ""
                for run in para.runs:
                    text = run.text
                    if run.bold: text = f"**{text}**"
                    if run.italic: text = f"*{text}*"
                    if run.underline: text = f"__{text}__"
                    markdown_line += text
                
                structured_markdown.append(f"[[ID:{para_id}]]\n{markdown_line}")
                plain_text_parts.append(para.text)
                para_idx_counter_body += 1

        elif isinstance(element, CT_Tbl):
            table = Table(element, doc)
            table_id_prefix = f"table_{table_idx_counter_body}"
            structured_markdown.append(f"[[TABLE_START:{table_id_prefix}]]")
            plain_text_parts.append(f"[جدول {table_idx_counter_body+1}]")

            # Convert table to markdown table format
            md_table = []
            for r_idx, row in enumerate(table.rows):
                row_text_parts = []
                row_plain_parts = []
                for c_idx, cell in enumerate(row.cells):
                    cell_id_prefix = f"{table_id_prefix}_r{r_idx}_c{c_idx}"
                    cell_para_idx_counter = 0
                    cell_markdown_content = ""
                    cell_plain_content = ""

                    for para_in_cell in cell.paragraphs:
                        if para_in_cell.text.strip():
                            cell_para_id = f"{cell_id_prefix}_p{cell_para_idx_counter}"
                            
                            md_line_cell = ""
                            for run in para_in_cell.runs:
                                text = run.text
                                if run.bold: text = f"**{text}**"
                                if run.italic: text = f"*{text}*"
                                if run.underline: text = f"__{text}__"
                                md_line_cell += text

                            # Add ID only to the first part of the cell content for clarity
                            if cell_para_idx_counter == 0:
                                cell_markdown_content += f"[[ID:{cell_para_id}]] {md_line_cell}"
                            else:
                                cell_markdown_content += f"\n[[ID:{cell_para_id}]] {md_line_cell}"
                            
                            cell_plain_content += para_in_cell.text + "\n"
                            cell_para_idx_counter += 1
                    
                    row_text_parts.append(cell_markdown_content.replace("\n", "<br>"))
                    row_plain_parts.append(cell_plain_content.strip())

                md_table.append("| " + " | ".join(row_text_parts) + " |")
                plain_text_parts.append(" | ".join(row_plain_parts))

                if r_idx == 0:
                    md_table.append("|" + " --- |" * len(row.cells))
            
            structured_markdown.extend(md_table)
            structured_markdown.append(f"[[TABLE_END:{table_id_prefix}]]")
            table_idx_counter_body += 1

    return "\n\n".join(structured_markdown), "\n".join(plain_text_parts)


def set_cell_direction_rtl(cell: _Cell):
    """Sets the visual direction of a table cell to RTL."""
    tcPr = cell._tc.get_or_add_tcPr() 
    bidiVisual = tcPr.find(qn('w:bidiVisual'))
    if bidiVisual is None:
        bidiVisual = OxmlElement('w:bidiVisual')
        tcPr.append(bidiVisual)


def _parse_markdown_to_parts_for_runs(text_line: str) -> list[dict]:
    """Helper to parse a line of markdown text for bold, italic, and underline into parts for runs."""
    parts_raw = re.split(r'(\*\*|\*|__)', text_line)
    
    parts = []
    is_bold = False
    is_italic = False
    is_underline = False
    
    for part in parts_raw:
        if part == '**': is_bold = not is_bold; continue
        if part == '*': is_italic = not is_italic; continue
        if part == '__': is_underline = not is_underline; continue
        
        if part:
            parts.append({
                "text": part,
                "bold": is_bold,
                "italic": is_italic,
                "underline": is_underline
            })
    return parts


def _add_paragraph_with_markdown_formatting(
    doc_or_cell, 
    style_name: str,
    text_content: str,
    contract_language: str,
    chosen_font: str,
    text_color: RGBColor | None = None,
    strike: bool = False, 
    is_list_item: bool = False,
    list_indent: Inches | None = None,
    first_line_indent_list: Inches | None = None 
):
    """
    Adds a new paragraph with specified text, parsing markdown for bold, italic, and underline.
    """
    if hasattr(doc_or_cell, 'add_paragraph'):
        p = doc_or_cell.add_paragraph(style=style_name)
    else: 
        if doc_or_cell.paragraphs:
            p = doc_or_cell.paragraphs[0]
            for run in list(p.runs): 
                p_element = run._element.getparent()
                if p_element is not None:
                    p_element.remove(run._element)
        else:
            p = doc_or_cell.add_paragraph()
        p.style = doc_or_cell.part.document.styles[style_name]

    if contract_language == 'ar':
        if style_name not in ['TitleStyle', 'BasmalaStyle', 'Heading2Style', 'Heading3Style']:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT 
        p.paragraph_format.rtl = True 
        if is_list_item and list_indent is not None:
            p.paragraph_format.left_indent = list_indent 
            if first_line_indent_list is not None: 
                 p.paragraph_format.first_line_indent = first_line_indent_list
    else: 
        if style_name not in ['TitleStyle', 'BasmalaStyle', 'Heading2Style', 'Heading3Style']:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.rtl = False 
        if is_list_item and list_indent is not None:
            p.paragraph_format.left_indent = list_indent

    parts = _parse_markdown_to_parts_for_runs(text_content)

    for part_info in parts:
        run = p.add_run(part_info["text"])
        if part_info["bold"]: run.bold = True
        if part_info["italic"]: run.italic = True
        if part_info["underline"]: run.underline = True
        run.font.rtl = (contract_language == 'ar') 
        run.font.name = chosen_font
        
        style_font_size = p.style.font.size if p.style and p.style.font else None
        run.font.size = style_font_size if style_font_size else Pt(12) 

        if text_color:
            run.font.color.rgb = text_color
        elif p.style and p.style.font and p.style.font.color and p.style.font.color.rgb:
            run.font.color.rgb = p.style.font.color.rgb
        
        if strike and (not text_color or text_color.rgb != RGBColor(255,0,0).rgb):
            run.font.strike = True
    return p


def _determine_style_and_text(line: str, contract_language: str) -> tuple[str, str, bool, bool]:
    """Helper to determine paragraph style and clean text from a markdown line."""
    current_style_name = 'NormalStyle'
    is_main_title = False
    is_list_item_flag = False
    text_for_paragraph_content = line.strip() 

    if line.strip() == "بسم الله الرحمن الرحيم":
        current_style_name = 'BasmalaStyle'
        text_for_paragraph_content = line.strip()
        return current_style_name, text_for_paragraph_content, is_main_title, is_list_item_flag

    if line.startswith('# '): 
        current_style_name = 'TitleStyle'
        is_main_title = True
        text_for_paragraph_content = re.sub(r'^#\s*', '', line).strip()
    elif contract_language == 'ar' and (
        re.match(r'^(البند)\s+(الأول|الثاني|الثالث|الرابع|الخامس|السادس|السابع|الثامن|التاسع|العاشر|الحادي عشر|الثاني عشر|الأخير|التمهيدي)\s*[:]?\s*$', line.strip()) or
        re.match(r'^(المادة)\s+\d+\s*[:]?\s*$', line.strip()) 
    ):
        current_style_name = 'Heading2Style'
        text_for_paragraph_content = line.strip() 
    elif contract_language == 'en' and (
        re.match(r'^(Clause|Article|Section)\s+\d+\s*[:]?\s*$', line.strip(), re.IGNORECASE) or
        re.match(r'^(Preamble|Preliminary Clause)\s*[:]?\s*$', line.strip(), re.IGNORECASE)
    ):
        current_style_name = 'Heading2Style'
        text_for_paragraph_content = line.strip()
    elif contract_language == 'ar' and (
        re.match(r'^(أولاً|ثانياً|ثالثاً|رابعاً|خامساً|سادساً|سابعاً|ثامناً|تاسعاً|عاشراً)\s*[:]', line.strip()) or
        re.match(r'^[أ-ي]\.\s+', line.strip()) 
    ):
        current_style_name = 'Heading3Style'
        text_for_paragraph_content = line.strip() 
    elif contract_language == 'en' and (
        re.match(r'^(Firstly|Secondly|Thirdly|Fourthly|Fifthly)\s*[:]', line.strip(), re.IGNORECASE) or
        re.match(r'^[A-Z]\.\s+', line.strip()) 
    ):
        current_style_name = 'Heading3Style'
        text_for_paragraph_content = line.strip()
    elif line.startswith('## '): 
        current_style_name = 'Heading2Style' 
        text_for_paragraph_content = re.sub(r'^##\s*', '', line).strip()
    elif line.startswith('### '): 
        current_style_name = 'Heading3Style'
        text_for_paragraph_content = re.sub(r'^###\s*', '', line).strip()
    elif line.startswith(("* ", "- ", "+ ")) or re.match(r'^\d+\.\s+', line):
        current_style_name = 'ListBulletStyle'
        is_list_item_flag = True
        text_for_paragraph_content = re.sub(r'^\s*[\*\-\+]+\s*|^\s*\d+\.\s*', '', line).strip()
    else:
        text_for_paragraph_content = line.strip() 
        
    text_for_paragraph_content = re.sub(r'^\[\[ID:.*?\]\]\s*', '', text_for_paragraph_content).strip()
    return current_style_name, text_for_paragraph_content, is_main_title, is_list_item_flag


def create_docx_from_llm_markdown(
    original_markdown_text: str, 
    output_path: str, 
    contract_language: str = 'ar', 
    terms_for_marking: list[dict] | dict | None = None,
    confirmed_modifications: dict | None = None
):
    """
    Creates a professional DOCX document from markdown text with term highlighting.
    Matches OldStrcturePerfectProject/doc_processing.py exactly.
    """
    try:
        doc = DocxDocument()
        chosen_font = "Arial" 
        
        for section in doc.sections:
            section.page_width = Inches(8.27) 
            section.page_height = Inches(11.69)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)

        styles = doc.styles
        basmala_style = styles.add_style('BasmalaStyle', WD_STYLE_TYPE.PARAGRAPH)
        basmala_format = basmala_style.paragraph_format
        basmala_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        basmala_format.space_before = Pt(0) 
        basmala_format.space_after = Pt(12) 
        basmala_font = basmala_style.font
        basmala_font.rtl = True 
        basmala_font.name = chosen_font 
        basmala_font.size = Pt(18) 
        basmala_font.bold = True

        title_style = styles.add_style('TitleStyle', WD_STYLE_TYPE.PARAGRAPH)
        title_format = title_style.paragraph_format
        title_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_format.space_after = Pt(18) 
        title_font = title_style.font
        title_font.rtl = (contract_language == 'ar')
        title_font.name = chosen_font
        title_font.size = Pt(20) 
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 0, 0)

        heading2_style = styles.add_style('Heading2Style', WD_STYLE_TYPE.PARAGRAPH)
        heading2_format = heading2_style.paragraph_format
        heading2_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if contract_language == 'ar' else WD_PARAGRAPH_ALIGNMENT.LEFT 
        heading2_format.space_before = Pt(12)
        heading2_format.space_after = Pt(6)
        heading2_font = heading2_style.font
        heading2_font.rtl = (contract_language == 'ar')
        heading2_font.name = chosen_font
        heading2_font.size = Pt(16) 
        heading2_font.bold = True
        heading2_font.underline = False 

        heading3_style = styles.add_style('Heading3Style', WD_STYLE_TYPE.PARAGRAPH)
        heading3_format = heading3_style.paragraph_format
        heading3_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT if contract_language == 'ar' else WD_PARAGRAPH_ALIGNMENT.LEFT
        heading3_format.space_before = Pt(10)
        heading3_format.space_after = Pt(4)
        heading3_font = heading3_style.font
        heading3_font.rtl = (contract_language == 'ar')
        heading3_font.name = chosen_font
        heading3_font.size = Pt(14) 
        heading3_font.bold = True 
        heading3_font.underline = False 

        normal_style = styles.add_style('NormalStyle', WD_STYLE_TYPE.PARAGRAPH)
        normal_format = normal_style.paragraph_format
        normal_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT if contract_language == 'ar' else WD_PARAGRAPH_ALIGNMENT.LEFT
        if contract_language == 'ar': 
            normal_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY_LOW 
        normal_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        normal_format.space_after = Pt(6) 
        normal_font = normal_style.font
        normal_font.rtl = (contract_language == 'ar')
        normal_font.name = chosen_font
        normal_font.size = Pt(12) 
        
        list_indent_val = Inches(0.5 if contract_language == 'ar' else 0.25)
        first_line_indent_val_list = Inches(-0.25) if contract_language == 'ar' else Inches(0) 
        
        list_style = styles.add_style('ListBulletStyle', WD_STYLE_TYPE.PARAGRAPH)
        list_format = list_style.paragraph_format
        list_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT if contract_language == 'ar' else WD_PARAGRAPH_ALIGNMENT.LEFT
        list_format.left_indent = list_indent_val 
        if contract_language == 'ar':
            list_format.first_line_indent = first_line_indent_val_list 
        list_format.space_after = Pt(4)
        list_font = list_style.font
        list_font.rtl = (contract_language == 'ar')
        list_font.name = chosen_font
        list_font.size = Pt(12)
        
        table_style = doc.styles.add_style('CustomTable', WD_STYLE_TYPE.TABLE)
        table_style.font.name = chosen_font
        table_style.font.size = Pt(10)

        lines = original_markdown_text.split('\n')
        processed_markdown_text = original_markdown_text
        if lines and lines[0].strip() == "بسم الله الرحمن الرحيم":
            _add_paragraph_with_markdown_formatting(doc, 'BasmalaStyle', lines[0].strip(), 'ar', chosen_font) 
            processed_markdown_text = "\n".join(lines[1:]) 
        
        if isinstance(terms_for_marking, list) and terms_for_marking: 
            logger.info(f"DOC_PROCESSING: Using new list-based term marking logic for PDF/TXT. {len(terms_for_marking)} terms.")
            current_markdown_pos = 0
            
            for term_idx, term_data in enumerate(terms_for_marking):
                term_text_original = term_data.get("term_text", "") 
                if not term_text_original.strip(): 
                    continue
                
                search_start_pos = current_markdown_pos
                
                def normalize_text(text):
                    """Normalize text for matching - remove extra whitespace."""
                    return ' '.join(text.split())
                
                def find_term_flexible(text, search_text, start_pos):
                    """Find term with flexible whitespace matching."""
                    normalized_search = normalize_text(search_text)
                    search_words = normalized_search.split()[:10]
                    if not search_words:
                        return None
                    
                    first_words = ' '.join(search_words[:3]) if len(search_words) >= 3 else ' '.join(search_words)
                    
                    pos = start_pos
                    while pos < len(text):
                        idx = text.find(search_words[0], pos)
                        if idx == -1:
                            break
                        
                        chunk = text[idx:idx + len(search_text) + 200]
                        if normalize_text(chunk).startswith(normalized_search[:100]):
                            end_pos = idx + len(search_text)
                            for end_offset in range(-20, 50):
                                test_end = end_pos + end_offset
                                if test_end <= len(text):
                                    chunk_to_test = text[idx:test_end]
                                    if normalize_text(chunk_to_test) == normalized_search:
                                        return (idx, chunk_to_test)
                            return (idx, text[idx:idx + len(search_text)])
                        pos = idx + 1
                    
                    return None
                
                match = None
                try:
                    term_text_pattern = re.escape(term_text_original).replace(r'\\n', r'\s*\\n\s*') 
                    term_text_pattern = term_text_pattern.replace(r'\n', r'\s*\n\s*')
                    term_text_pattern = term_text_pattern.replace(r'\ ', r'\s+')
                    match = re.search(term_text_pattern, processed_markdown_text[search_start_pos:], re.DOTALL)
                except re.error as re_err:
                    logger.warning(f"Regex error for term '{term_data.get('term_id')}': {re_err}. Trying flexible match.")
                    match = None
                
                if not match:
                    flexible_result = find_term_flexible(processed_markdown_text, term_text_original, search_start_pos)
                    if flexible_result:
                        found_pos, matched_text = flexible_result
                        class FlexibleMatch:
                            def __init__(self, pos, text, start_pos):
                                self._pos = pos - start_pos
                                self._text = text
                            def start(self): return self._pos
                            def group(self, _=0): return self._text
                        match = FlexibleMatch(found_pos, matched_text, search_start_pos)
                    else:
                        _found_pos = processed_markdown_text.find(term_text_original[:50], search_start_pos)
                        if _found_pos != -1:
                            actual_text = processed_markdown_text[_found_pos:_found_pos + len(term_text_original)]
                            class FallbackMatch:
                                def __init__(self, pos, text, start_pos):
                                    self._pos = pos - start_pos
                                    self._text = text
                                def start(self): return self._pos
                                def group(self, _=0): return self._text
                            match = FallbackMatch(_found_pos, actual_text, search_start_pos)

                if match:
                    found_pos_relative = match.start()
                    found_pos_absolute = search_start_pos + found_pos_relative
                    matched_text_in_doc = match.group(0) 

                    inter_term_text = processed_markdown_text[current_markdown_pos:found_pos_absolute]
                    if inter_term_text.strip():
                        for line_in_inter_term in inter_term_text.splitlines(): 
                             if "[[TABLE_" in line_in_inter_term: 
                                 continue
                             l_style, l_text, _, l_is_list = _determine_style_and_text(line_in_inter_term, contract_language)
                             _add_paragraph_with_markdown_formatting(doc, l_style, l_text, contract_language, chosen_font, is_list_item=l_is_list, list_indent=list_indent_val if l_is_list else None, first_line_indent_list=first_line_indent_val_list if l_is_list and contract_language == 'ar' else None)
                    
                    logger.info(f"Found term '{term_data.get('term_id')}' at pos {found_pos_absolute}")
                    initial_is_valid = term_data.get("is_valid_sharia", True)
                    is_confirmed = term_data.get("is_confirmed_by_user", False)
                    confirmed_text_content = term_data.get("confirmed_modified_text") 
                    
                    term_lines_to_render_original = matched_text_in_doc.splitlines() 

                    if is_confirmed and confirmed_text_content and \
                       not initial_is_valid and confirmed_text_content.strip() != term_text_original.strip():
                        logger.info(f"Applying MARKING: Red original, Green new for term {term_data.get('term_id')}")
                        for line_in_term in term_lines_to_render_original:
                             if "[[TABLE_" in line_in_term: 
                                 continue
                             l_style, l_text, _, l_is_list = _determine_style_and_text(line_in_term, contract_language)
                             _add_paragraph_with_markdown_formatting(doc, l_style, l_text, contract_language, chosen_font, text_color=RGBColor(255,0,0), strike=False, is_list_item=l_is_list, list_indent=list_indent_val if l_is_list else None, first_line_indent_list=first_line_indent_val_list if l_is_list and contract_language == 'ar' else None) 
                        
                        sep_para = doc.add_paragraph(style='NormalStyle')
                        sep_para.paragraph_format.rtl = (contract_language == 'ar')
                        sep_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT if contract_language == 'ar' else WD_PARAGRAPH_ALIGNMENT.LEFT
                        sep_run = sep_para.add_run(("التعديل المؤكد: " if contract_language == 'ar' else "Confirmed Modification: "))
                        sep_run.font.size = Pt(10)
                        sep_run.italic = True
                        sep_run.font.name = chosen_font
                        sep_run.font.rtl = (contract_language == 'ar')
                        
                        for line_in_confirmed_text in confirmed_text_content.splitlines():
                            if "[[TABLE_" in line_in_confirmed_text: 
                                continue
                            l_style, l_text, _, l_is_list = _determine_style_and_text(line_in_confirmed_text, contract_language)
                            _add_paragraph_with_markdown_formatting(doc, l_style, l_text, contract_language, chosen_font, text_color=RGBColor(0,128,0), is_list_item=l_is_list, list_indent=list_indent_val if l_is_list else None, first_line_indent_list=first_line_indent_val_list if l_is_list and contract_language == 'ar' else None)
                    else:
                        text_to_render_for_term = matched_text_in_doc 
                        final_text_color_for_term = None
                        if is_confirmed and confirmed_text_content:
                            text_to_render_for_term = confirmed_text_content 
                            final_text_color_for_term = RGBColor(0,128,0) 
                            logger.info(f"Applying MARKING: Green (confirmed) for term {term_data.get('term_id')}")
                        elif not initial_is_valid:
                            final_text_color_for_term = RGBColor(255,0,0) 
                            logger.info(f"Applying MARKING: Red (initially invalid) for term {term_data.get('term_id')}")
                        
                        for line_in_term_render in text_to_render_for_term.splitlines():
                            if "[[TABLE_" in line_in_term_render: 
                                continue
                            l_style, l_text, _, l_is_list = _determine_style_and_text(line_in_term_render, contract_language)
                            _add_paragraph_with_markdown_formatting(doc, l_style, l_text, contract_language, chosen_font, text_color=final_text_color_for_term, strike=False, is_list_item=l_is_list, list_indent=list_indent_val if l_is_list else None, first_line_indent_list=first_line_indent_val_list if l_is_list and contract_language == 'ar' else None) 
                    
                    current_markdown_pos = found_pos_absolute + len(matched_text_in_doc)
                else:
                    logger.warning(f"Term '{term_data.get('term_id')}' text not found sequentially from pos {search_start_pos}")
            
            if current_markdown_pos < len(processed_markdown_text):
                remaining_text = processed_markdown_text[current_markdown_pos:]
                for line_in_remaining in remaining_text.splitlines():
                    if "[[TABLE_" in line_in_remaining: 
                        continue
                    l_style, l_text, _, l_is_list = _determine_style_and_text(line_in_remaining, contract_language)
                    _add_paragraph_with_markdown_formatting(doc, l_style, l_text, contract_language, chosen_font, is_list_item=l_is_list, list_indent=list_indent_val if l_is_list else None, first_line_indent_list=first_line_indent_val_list if l_is_list and contract_language == 'ar' else None)
        else:
            logger.info(f"DOC_PROCESSING: Using old dict-based or no-term marking logic. terms_for_marking type: {type(terms_for_marking)}")
            lines_to_process = processed_markdown_text.split('\n')
            i = 0
            while i < len(lines_to_process):
                line = lines_to_process[i].strip()
                if not line or "[[TABLE_" in line: 
                    i += 1
                    continue

                if line.startswith('|') and line.endswith('|') and line.count('|') > 1:
                    table_lines = []
                    temp_i = i
                    while temp_i < len(lines_to_process) and lines_to_process[temp_i].strip().startswith('|') and lines_to_process[temp_i].strip().endswith('|'):
                        table_lines.append(lines_to_process[temp_i].strip())
                        temp_i += 1
                    if len(table_lines) > 1 and re.match(r'\|(\s*:?-+:?\s*\|)+', table_lines[1]): 
                        header_row_content = [h.strip() for h in table_lines[0].strip('|').split('|')]
                        num_cols = len(header_row_content)
                        if num_cols > 0:
                            table_data_rows = []
                            for row_line_idx in range(2, len(table_lines)):
                                row_content_raw = [cell.strip().replace('<br>', '\n') for cell in table_lines[row_line_idx].strip('|').split('|')]
                                row_content = row_content_raw + [''] * (num_cols - len(row_content_raw)) if len(row_content_raw) < num_cols else row_content_raw[:num_cols]
                                table_data_rows.append(row_content)
                            if table_data_rows:
                                doc_table = doc.add_table(rows=1, cols=num_cols)
                                doc_table.style = 'CustomTable'
                                if contract_language == 'ar':
                                    doc_table.table_direction = WD_TABLE_DIRECTION.RTL
                                    doc_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
                                else:
                                    doc_table.table_direction = WD_TABLE_DIRECTION.LTR
                                    doc_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                                hdr_cells = doc_table.rows[0].cells
                                for col_idx, header_text in enumerate(header_row_content):
                                    cell_p = hdr_cells[col_idx].paragraphs[0]
                                    cell_p.text = ""
                                    _add_paragraph_with_markdown_formatting(hdr_cells[col_idx], 'Normal', re.sub(r'\[\[ID:.*?\]\]\s*', '', header_text).strip(), contract_language, chosen_font)
                                    hdr_cells[col_idx].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    if contract_language == 'ar':
                                        set_cell_direction_rtl(hdr_cells[col_idx])
                                for data_row_content in table_data_rows:
                                    row_cells = doc_table.add_row().cells
                                    for col_idx, cell_text in enumerate(data_row_content):
                                        cell_p = row_cells[col_idx].paragraphs[0]
                                        cell_p.text = ""
                                        _add_paragraph_with_markdown_formatting(row_cells[col_idx], 'Normal', re.sub(r'\[\[ID:.*?\]\]\s*', '', cell_text).strip(), contract_language, chosen_font)
                                        if contract_language == 'ar':
                                            set_cell_direction_rtl(row_cells[col_idx])
                                doc.add_paragraph()
                                i = temp_i
                                continue
                
                current_style_name, text_for_paragraph_content, is_main_title, is_list_item_flag = _determine_style_and_text(line, contract_language)
                
                term_status_info = None
                if isinstance(terms_for_marking, dict): 
                    clean_para_text_for_match = re.sub(r'^\[\[ID:.*?\]\]\s*', '', text_for_paragraph_content).strip()
                    term_status_info = terms_for_marking.get(clean_para_text_for_match) 

                if term_status_info: 
                    is_confirmed = term_status_info.get("is_confirmed", False)
                    confirmed_text_content = term_status_info.get("confirmed_text")
                    initial_is_valid = term_status_info.get("initial_is_valid", True)
                    current_original_text_for_term = clean_para_text_for_match 

                    if is_confirmed and confirmed_text_content and \
                       not initial_is_valid and confirmed_text_content.strip() != current_original_text_for_term.strip():
                        _add_paragraph_with_markdown_formatting(doc, current_style_name, current_original_text_for_term, contract_language, chosen_font, text_color=RGBColor(255,0,0), strike=False, is_list_item=is_list_item_flag, list_indent=list_indent_val if is_list_item_flag else None, first_line_indent_list=first_line_indent_val_list if is_list_item_flag and contract_language == 'ar' else None)
                        sep_para = doc.add_paragraph(style='NormalStyle')
                        sep_para.paragraph_format.rtl = (contract_language == 'ar')
                        sep_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT if contract_language == 'ar' else WD_PARAGRAPH_ALIGNMENT.LEFT
                        sep_run = sep_para.add_run(("التعديل المؤكد: " if contract_language == 'ar' else "Confirmed Modification: "))
                        sep_run.font.size = Pt(10)
                        sep_run.italic = True
                        sep_run.font.name = chosen_font
                        sep_run.font.rtl = (contract_language == 'ar')
                        _add_paragraph_with_markdown_formatting(doc, current_style_name, confirmed_text_content, contract_language, chosen_font, text_color=RGBColor(0,128,0), is_list_item=is_list_item_flag, list_indent=list_indent_val if is_list_item_flag else None, first_line_indent_list=first_line_indent_val_list if is_list_item_flag and contract_language == 'ar' else None)
                    else:
                        text_to_render = current_original_text_for_term
                        final_text_color = None
                        if is_confirmed and confirmed_text_content:
                            text_to_render = confirmed_text_content
                            final_text_color = RGBColor(0,128,0)
                        elif not initial_is_valid:
                            final_text_color = RGBColor(255,0,0)
                        _add_paragraph_with_markdown_formatting(doc, current_style_name, text_to_render, contract_language, chosen_font, text_color=final_text_color, strike=False, is_list_item=is_list_item_flag, list_indent=list_indent_val if is_list_item_flag else None, first_line_indent_list=first_line_indent_val_list if is_list_item_flag and contract_language == 'ar' else None)
                else: 
                    _add_paragraph_with_markdown_formatting(doc, current_style_name, text_for_paragraph_content, contract_language, chosen_font, text_color=None, strike=False, is_list_item=is_list_item_flag, list_indent=list_indent_val if is_list_item_flag else None, first_line_indent_list=first_line_indent_val_list if is_list_item_flag and contract_language == 'ar' else None)
                
                i += 1
        
        signature_found = any(sig_ar in line_text or sig_en in line_text 
                              for line_text in processed_markdown_text.split('\n') 
                              for sig_ar in ["وحرر هذا العقد", "التوقيعات", "الطرف الأول", "الطرف الثاني", "الشاهد الأول", "الشاهد الثاني"] 
                              for sig_en in ["This contract was made", "Signatures", "Party One", "Party Two", "First Witness", "Second Witness"])
        
        if not signature_found:
            doc.add_paragraph() 
            if contract_language == 'ar':
                p_sig_text = doc.add_paragraph(style='NormalStyle')
                p_sig_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p_sig_text.add_run("وحرر هذا العقد من نسختين بيد كل طرف نسخة للعمل بموجبها عند اللزوم.").font.name = chosen_font
                
                doc.add_paragraph() 
                
                sig_heading = doc.add_paragraph(style='Heading3Style')
                sig_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                sig_heading.add_run("التوقيعات").font.name = chosen_font
                
                table_sig = doc.add_table(rows=1, cols=2)
                table_sig.style = 'CustomTable'
                table_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
                if contract_language == 'ar':
                    table_sig.table_direction = WD_TABLE_DIRECTION.RTL

                def add_sig_cell_content(cell, party_name_text):
                    p_party_name = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                    p_party_name.text = "" 
                    run_party_name = p_party_name.add_run(party_name_text)
                    run_party_name.font.name = chosen_font
                    run_party_name.font.bold = True
                    run_party_name.font.size = Pt(12)
                    p_party_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    if contract_language == 'ar':
                        p_party_name.paragraph_format.rtl = True
                    
                    cell.add_paragraph(f"الإسم: \t\t\t", style='NormalStyle').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT if contract_language == 'ar' else WD_PARAGRAPH_ALIGNMENT.LEFT
                    cell.add_paragraph(f"بطاقة رقم قومي: \t\t", style='NormalStyle').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT if contract_language == 'ar' else WD_PARAGRAPH_ALIGNMENT.LEFT
                    cell.add_paragraph(f"التوقيع: \t\t\t", style='NormalStyle').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT if contract_language == 'ar' else WD_PARAGRAPH_ALIGNMENT.LEFT
                    cell.add_paragraph("\n") 

                add_sig_cell_content(table_sig.cell(0, 1), "الطرف الأول (البائعة)") 
                add_sig_cell_content(table_sig.cell(0, 0), "الطرف الثاني (المشترية)") 
                
                doc.add_paragraph() 
                witness_heading = doc.add_paragraph(style='Heading3Style')
                witness_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                witness_heading.add_run("توقيع الشهود").font.name = chosen_font

                table_witness = doc.add_table(rows=1, cols=2)
                table_witness.style = 'CustomTable'
                table_witness.alignment = WD_TABLE_ALIGNMENT.CENTER
                if contract_language == 'ar':
                    table_witness.table_direction = WD_TABLE_DIRECTION.RTL
                
                add_sig_cell_content(table_witness.cell(0, 1), "الشاهد الأول")
                add_sig_cell_content(table_witness.cell(0, 0), "الشاهد الثاني")

            else: 
                p_sig = doc.add_paragraph("This contract is executed in two counterparts...", style='NormalStyle')
                p_sig.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                doc.add_paragraph("")
                signature_section = doc.add_paragraph("Signatures", style='Heading2Style')
                signature_section.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                table_sig = doc.add_table(rows=2, cols=2)
                table_sig.style = 'CustomTable'
                table_sig.table_direction = WD_TABLE_DIRECTION.LTR
                table_sig.alignment = WD_TABLE_ALIGNMENT.LEFT
                cell1_sig = table_sig.cell(0, 0)
                cell1_para_sig = cell1_sig.paragraphs[0]
                cell1_para_sig.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell1_run_sig = cell1_para_sig.add_run("Party One")
                cell1_run_sig.font.name = chosen_font
                cell1_run_sig.font.bold = True
                cell2_sig = table_sig.cell(0, 1)
                cell2_para_sig = cell2_sig.paragraphs[0]
                cell2_para_sig.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell2_run_sig = cell2_para_sig.add_run("Party Two")
                cell2_run_sig.font.name = chosen_font
                cell2_run_sig.font.bold = True
                table_sig.cell(1, 0).text = "\nName:\nID:\nSignature:\n____________________"
                table_sig.cell(1, 1).text = "\nName:\nID:\nSignature:\n____________________"

        ensure_dir(os.path.dirname(output_path))
        doc.save(output_path)
        
        logger.info(f"DOCX document created successfully: {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"Error creating DOCX document: {e}")
        traceback.print_exc()
        raise ValueError(f"فشل إنشاء DOCX: {e}")


def _find_libreoffice_path() -> str:
    """Find LibreOffice executable path - cross-platform."""
    try:
        libreoffice_path = current_app.config.get('LIBREOFFICE_PATH', '')
    except RuntimeError:
        libreoffice_path = ''
    
    if libreoffice_path:
        clean_path = libreoffice_path.strip()
        if clean_path.startswith('r"') and clean_path.endswith('"'):
            clean_path = clean_path[2:-1]
        elif clean_path.startswith('"') and clean_path.endswith('"'):
            clean_path = clean_path[1:-1]
        if os.path.exists(clean_path):
            return clean_path
    
    if os.name == 'nt':
        windows_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            os.path.expandvars(r"%PROGRAMFILES%\LibreOffice\program\soffice.exe"),
            os.path.expandvars(r"%PROGRAMFILES(X86)%\LibreOffice\program\soffice.exe"),
        ]
        for path in windows_paths:
            if os.path.exists(path):
                logger.info(f"Found LibreOffice at: {path}")
                return path
        return "soffice.exe"
    else:
        linux_paths = [
            "/usr/bin/libreoffice",
            "/usr/bin/soffice",
            "/usr/local/bin/libreoffice",
            "/usr/local/bin/soffice",
        ]
        for path in linux_paths:
            if os.path.exists(path):
                logger.info(f"Found LibreOffice at: {path}")
                return path
        return "libreoffice"


def convert_docx_to_pdf(docx_path: str, output_folder: str) -> str:
    """
    Converts a DOCX file to PDF using LibreOffice directly.
    Returns the path to the generated PDF, or raises an exception on failure.
    Automatically finds LibreOffice on Windows and Linux.
    """
    if not os.path.exists(docx_path):
        logger.error(f"DOCX file not found for PDF conversion: {docx_path}")
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    ensure_dir(output_folder)

    pdf_filename = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
    pdf_output_path = os.path.join(output_folder, pdf_filename)

    soffice_cmd = _find_libreoffice_path()

    command = [
        soffice_cmd,
        '--headless',
        '--convert-to', 'pdf',
        '--outdir', output_folder,
        docx_path
    ]

    logger.info(f"Attempting PDF conversion with command: {' '.join(command)}")

    try:
        is_windows = os.name == 'nt'
        startupinfo = None
        if is_windows:
            startupinfo = subprocess.STARTUPINFO()
            startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
            startupinfo.wShowWindow = subprocess.SW_HIDE

        result = subprocess.run(
            command,
            capture_output=True,
            text=True,
            check=False,
            timeout=180,
            startupinfo=startupinfo
        )

        if result.returncode != 0:
            logger.error(f"Error converting DOCX to PDF. LibreOffice/soffice process exited with code: {result.returncode}")
            logger.error(f"soffice stdout: {result.stdout}")
            logger.error(f"soffice stderr: {result.stderr}")
            if os.path.exists(pdf_output_path) and os.path.getsize(pdf_output_path) == 0:
                os.remove(pdf_output_path)
            raise Exception(f"LibreOffice/soffice conversion failed. STDERR: {result.stderr[:1000]}")

        if os.path.exists(pdf_output_path) and os.path.getsize(pdf_output_path) > 0:
            logger.info(f"PDF conversion successful: {pdf_output_path}")
            return pdf_output_path
        else:
            logger.error(f"PDF file not created or is empty at {pdf_output_path} despite successful soffice exit code.")
            logger.error(f"soffice stdout: {result.stdout}")
            logger.error(f"soffice stderr: {result.stderr}")
            if os.path.exists(pdf_output_path):
                os.remove(pdf_output_path)
            raise Exception("PDF file not created or is empty after LibreOffice/soffice execution.")

    except FileNotFoundError:
        logger.error(f"CRITICAL ERROR: '{soffice_cmd}' command not found. Please ensure LibreOffice is installed and '{soffice_cmd}' is in your system PATH.")
        raise Exception(f"PDF conversion tool ('{soffice_cmd}') not found. Check LibreOffice installation and PATH/config.")
    except subprocess.TimeoutExpired:
        logger.error(f"PDF conversion timed out for {docx_path}.")
        if os.path.exists(pdf_output_path):
            os.remove(pdf_output_path)
        raise Exception("PDF conversion timed out.")
    except Exception as e:
        logger.error(f"An unexpected error occurred during PDF conversion for {docx_path}: {e}")
        traceback.print_exc()
        if os.path.exists(pdf_output_path):
            os.remove(pdf_output_path)
        raise Exception(f"PDF conversion failed: {str(e)}")
